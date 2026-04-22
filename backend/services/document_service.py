import hashlib
import io
import logging
from pathlib import Path
from datetime import datetime
from uuid import uuid4
from sqlalchemy.orm import Session

from models import Document, DocumentChunk
from config import PROJECTS_DIR
from services import object_storage

logger = logging.getLogger("reco.docs")


def _pg_safe_text(s: str) -> str:
    """Postgres rejects NUL in text / JSON string values."""
    return (s or "").replace("\x00", "")


def _try_write_local(project_id: str, filename: str, content: bytes) -> None:
    """Best-effort local-disk copy when Supabase storage isn't writing.

    Pure side-effect: failures are logged and swallowed because the canonical
    record (Document + chunks) lives in Postgres, and Drive bytes can be re-fetched.
    """
    try:
        doc_dir = PROJECTS_DIR / project_id / "docs"
        doc_dir.mkdir(parents=True, exist_ok=True)
        safe_name = Path(filename).name or "file"
        (doc_dir / safe_name).write_bytes(content)
    except OSError as exc:
        logger.warning("Local doc write failed for %s/%s (%s); continuing.",
                       project_id, filename, exc)


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _extract_text(content: bytes, mime_type: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()

    if mime_type == "application/pdf" or ext == ".pdf":
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            out = io.StringIO()
            extract_text_to_fp(io.BytesIO(content), out, laparams=LAParams())
            return _pg_safe_text(out.getvalue())
        except Exception:
            return ""

    if (mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return _pg_safe_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()))
        except Exception:
            return ""

    if mime_type and mime_type.startswith("text/"):
        return _pg_safe_text(content.decode("utf-8", errors="replace"))

    return ""


def _chunk_text(text: str, chunk_tokens: int = 1000) -> list[str]:
    """Split text into ~chunk_tokens-sized chunks at sentence boundaries."""
    # Approx 4 chars per token
    chunk_chars = chunk_tokens * 4
    sentences = text.replace("\n", " \n ").split(". ")
    chunks, current = [], ""
    for sent in sentences:
        if len(current) + len(sent) > chunk_chars and current:
            chunks.append(current.strip())
            current = sent + ". "
        else:
            current += sent + ". "
    if current.strip():
        chunks.append(current.strip())
    return [_pg_safe_text(c) for c in chunks if c]


def process_upload(project_id: str, filename: str, content: bytes,
                   mime_type: str, db: Session) -> Document:
    filename = _pg_safe_text(filename) or "file"
    file_hash = _sha256(content)

    # Dedup: skip if same file already in this project
    existing = db.query(Document).filter_by(project_id=project_id, file_hash=file_hash).first()
    if existing:
        return existing

    doc_id = str(uuid4())
    storage_key = object_storage.save_file(
        project_id, doc_id, filename, content, mime_type or None,
    )
    if storage_key is None:
        _try_write_local(project_id, filename, content)

    text = _extract_text(content, mime_type, filename)
    chunks_text = _chunk_text(text)

    doc = Document(
        id=doc_id,
        project_id=project_id,
        filename=filename,
        source="upload",
        mime_type=mime_type,
        size_bytes=len(content),
        file_hash=file_hash,
        storage_object_key=storage_key,
    )
    db.add(doc)
    db.flush()

    for i, chunk in enumerate(chunks_text):
        token_count = len(chunk) // 4
        db.add(DocumentChunk(document_id=doc.id, chunk_index=i, text=chunk, token_count=token_count))

    db.commit()
    db.refresh(doc)
    return doc


def process_bytes(project_id: str, filename: str, content: bytes,
                  mime_type: str, source: str, db: Session,
                  drive_file_id: str = None, gmail_message_id: str = None) -> Document | None:
    """Generic ingestion used by Drive and Gmail services."""
    filename = _pg_safe_text(filename) or "file"
    file_hash = _sha256(content)
    existing = db.query(Document).filter_by(project_id=project_id, file_hash=file_hash).first()
    if existing:
        # Same bytes already ingested (e.g. uploaded or via Gmail). Link this Drive file id when
        # missing so Drive sync does not report a false "skip"; do not steal the link from a
        # different Drive file that happens to share the same hash.
        if drive_file_id:
            if existing.drive_file_id in (None, drive_file_id):
                if existing.drive_file_id != drive_file_id:
                    existing.drive_file_id = drive_file_id
                    db.commit()
                    db.refresh(existing)
                return existing
            return None
        return None  # duplicate (non-Drive ingest)

    doc_id = str(uuid4())
    storage_key = object_storage.save_file(
        project_id, doc_id, filename, content, mime_type or None,
    )
    if storage_key is None:
        _try_write_local(project_id, filename, content)

    text = _extract_text(content, mime_type, filename)
    chunks_text = _chunk_text(text)

    doc = Document(
        id=doc_id,
        project_id=project_id,
        filename=filename,
        source=source,
        mime_type=mime_type,
        size_bytes=len(content),
        file_hash=file_hash,
        drive_file_id=drive_file_id,
        gmail_message_id=gmail_message_id,
        storage_object_key=storage_key,
    )
    db.add(doc)
    db.flush()

    for i, chunk in enumerate(chunks_text):
        db.add(DocumentChunk(document_id=doc.id, chunk_index=i, text=chunk,
                             token_count=len(chunk) // 4))

    db.commit()
    db.refresh(doc)
    return doc
